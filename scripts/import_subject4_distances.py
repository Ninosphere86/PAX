#!/usr/bin/env python3

import argparse
import json
import re
from pathlib import Path

from docx import Document


QUESTION_RE = re.compile(r"^N02_(\d{2,3})\s+(.*)$")
OPTION_RE = re.compile(r"([ABCD])、(.*?)(?=\s+[ABCD]、|$)")
ANSWER_RE = re.compile(r"^【答案】([A-D]+)\s+【(单选题|多选题)】$")

EXPLANATION_FIXES = {
    13: "无人看守铁路道口前100米，注意提前控制车速",
    15: "事故易发路段，要谨慎驾驶",
    20: "地点距离标志用于了解前方重要地点和距离",
    23: "图中为右侧出口预告标志，箭头指向右上方",
}


def parse_answer_table(document):
    if not document.tables:
        raise ValueError("文档中缺少答案速查表")

    answers = {}
    for row in document.tables[-1].rows[1:]:
        values = [cell.text.strip() for cell in row.cells]
        for index in range(0, len(values), 2):
            number = values[index]
            answer = values[index + 1] if index + 1 < len(values) else ""
            if number:
                answers[int(number)] = answer
    return answers


def parse_questions(document):
    questions = []
    current = None

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        question_match = QUESTION_RE.match(text)
        if question_match:
            if current:
                questions.append(current)
            current = {
                "number": int(question_match.group(1)),
                "title": question_match.group(2).strip(),
                "options": {},
                "answer": "",
                "type": "",
                "explanation": "",
                "detailedExplanation": "",
            }
            continue

        if not current or not text:
            continue

        answer_match = ANSWER_RE.match(text)
        if answer_match:
            current["answer"] = answer_match.group(1)
            current["type"] = answer_match.group(2)
        elif text.startswith("【简析】"):
            current["explanation"] = text.removeprefix("【简析】").strip()
        elif text.startswith("【解析】"):
            current["detailedExplanation"] = text.removeprefix("【解析】").strip()
        elif text.startswith("A、"):
            current["options"] = {
                key: value.strip() for key, value in OPTION_RE.findall(text)
            }

    if current:
        questions.append(current)
    return questions


def validate(questions, answer_table):
    errors = []
    expected = list(range(1, 47))
    actual = [question["number"] for question in questions]
    if actual != expected:
        errors.append(f"题号不连续：expected 1..46, actual {actual}")

    for question in questions:
        number = question["number"]
        for field in ["title", "options", "answer", "type", "explanation", "detailedExplanation"]:
            if not question[field]:
                errors.append(f"第 {number} 题缺少 {field}")

        missing_answers = sorted(set(question["answer"]) - set(question["options"]))
        if missing_answers:
            errors.append(f"第 {number} 题答案不在选项中：{missing_answers}")
        if answer_table.get(number) != question["answer"]:
            errors.append(
                f"第 {number} 题正文答案 {question['answer']} 与速查表 {answer_table.get(number)} 不一致"
            )

    if len(answer_table) != 46:
        errors.append(f"答案速查表数量应为 46，实际为 {len(answer_table)}")
    if errors:
        raise ValueError("\n".join(errors))


def build_records(questions, image_dir):
    records = []
    for question in questions:
        number = question["number"]
        code = f"C02_{number:02d}"
        image_name = f"d02-{number:02d}-v2.webp"
        if not (image_dir / image_name).is_file():
            raise FileNotFoundError(f"缺少重制题图：{image_name}")

        records.append(
            {
                # Keep the historical ID stable so existing shared overrides still apply.
                "id": f"subject4-distances-D02_{number:02d}",
                "code": code,
                "title": question["title"],
                "image": f"/question-images/{image_name}",
                "options": {key: question["options"].get(key, "") for key in "ABCD"},
                "type": question["type"],
                "category": "科四场景类-C",
                "section": "C02",
                "difficulty": "基础",
                "answer": question["answer"],
                "explanation": EXPLANATION_FIXES.get(number, question["explanation"]),
                "detailedExplanation": question["detailedExplanation"].replace("雨、雪、雾灯特殊天气", "雨、雪、雾等特殊天气"),
                "tags": ["科四距离", "C02"],
                "status": "已发布",
                "updatedAt": "2026-07-22",
            }
        )
    return records


def main():
    parser = argparse.ArgumentParser(description="导入科目四距离题 Word 题库")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--project", type=Path, default=Path.cwd())
    parser.add_argument("--apply", action="store_true")
    args = parser.parse_args()

    document = Document(args.docx)
    questions = parse_questions(document)
    answer_table = parse_answer_table(document)
    validate(questions, answer_table)

    report = {
        "question_count": len(questions),
        "single_choice_count": sum(question["type"] == "单选题" for question in questions),
        "multiple_choice_count": sum(question["type"] == "多选题" for question in questions),
        "source_image_count": len(document.inline_shapes),
        "answer_table_count": len(answer_table),
        "applied": False,
    }

    if args.apply:
        project = args.project.resolve()
        bank_path = project / "app" / "question-bank.json"
        records = build_records(questions, project / "public" / "question-images")
        bank = json.loads(bank_path.read_text(encoding="utf-8"))
        bank = [
            item
            for item in bank
            if item.get("section") not in {"D02", "C02"}
            and not item.get("code", "").startswith(("D02_", "C02_"))
        ]
        bank.extend(records)
        bank_path.write_text(
            json.dumps(bank, ensure_ascii=False, separators=(",", ":")),
            encoding="utf-8",
        )
        report.update(
            {
                "applied": True,
                "written_question_count": len(records),
                "written_image_count": len(records),
                "total_bank_count": len(bank),
                "codes": [record["code"] for record in records],
            }
        )

    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
