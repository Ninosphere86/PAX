#!/usr/bin/env python3

import argparse
import json
import re
from pathlib import Path

from docx import Document


QUESTION_RE = re.compile(r"^L01_(\d{2,3})\s+(.*)$")
OPTION_RE = re.compile(r"([ABCD])、(.*?)(?=\s+[ABCD]、|$)")
ANSWER_RE = re.compile(r"^【答案】([A-D]+)\s+【(单选题|多选题)】$")


def parse_answer_table(document):
    if not document.tables:
        raise ValueError("文档中缺少答案速查表")

    answers = {}
    for row in document.tables[-1].rows[1:]:
        values = [cell.text.strip() for cell in row.cells]
        for index in range(0, len(values), 2):
            number = values[index]
            answer = values[index + 1] if index + 1 < len(values) else ""
            if number:
                answers[int(number)] = answer
    return answers


def parse_questions(document):
    questions = []
    current = None

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        question_match = QUESTION_RE.match(text)
        if question_match:
            if current:
                questions.append(current)
            current = {
                "number": int(question_match.group(1)),
                "title": question_match.group(2).strip(),
                "options": {},
                "answer": "",
                "type": "",
                "explanation": "",
                "detailedExplanation": "",
            }
            continue

        if not current or not text:
            continue

        answer_match = ANSWER_RE.match(text)
        if answer_match:
            current["answer"] = answer_match.group(1)
            current["type"] = answer_match.group(2)
        elif text.startswith("【简析】"):
            current["explanation"] = text.removeprefix("【简析】").strip()
        elif text.startswith("【解析】"):
            current["detailedExplanation"] = text.removeprefix("【解析】").strip()
        elif text.startswith("A、"):
            current["options"] = {
                key: value.strip() for key, value in OPTION_RE.findall(text)
            }

    if current:
        questions.append(current)
    return questions


def validate(questions, answer_table):
    errors = []
    actual = [question["number"] for question in questions]
    if actual != [1, 2, 3, 4]:
        errors.append(f"题号不连续：expected 1..4, actual {actual}")

    for question in questions:
        number = question["number"]
        for field in ["title", "options", "answer", "type", "explanation", "detailedExplanation"]:
            if not question[field]:
                errors.append(f"第 {number} 题缺少 {field}")

        missing_answers = sorted(set(question["answer"]) - set(question["options"]))
        if missing_answers:
            errors.append(f"第 {number} 题答案不在选项中：{missing_answers}")
        if answer_table.get(number) != question["answer"]:
            errors.append(
                f"第 {number} 题正文答案 {question['answer']} 与速查表 {answer_table.get(number)} 不一致"
            )

    if len(answer_table) != 4:
        errors.append(f"答案速查表数量应为 4，实际为 {len(answer_table)}")
    if errors:
        raise ValueError("\n".join(errors))


def build_records(questions, image_dir):
    records = []
    for question in questions:
        number = question["number"]
        code = f"C01_{number:02d}"
        image_name = f"a01-{number:02d}-v2.webp"
        if not (image_dir / image_name).is_file():
            raise FileNotFoundError(f"缺少重制题图：{image_name}")

        records.append(
            {
                # Keep the historical ID stable so existing shared overrides still apply.
                "id": f"subject4-drunk-driving-A01_{number:02d}",
                "code": code,
                "title": question["title"],
                "image": f"/question-images/{image_name}",
                "options": {key: question["options"].get(key, "") for key in "ABCD"},
                "type": question["type"],
                "category": "科四场景类-C",
                "section": "C01",
                "difficulty": "基础",
                "answer": question["answer"],
                "explanation": question["explanation"],
                "detailedExplanation": question["detailedExplanation"],
                "tags": ["科四酒驾", "C01"],
                "status": "已发布",
                "updatedAt": "2026-07-22",
            }
        )
    return records


def main():
    parser = argparse.ArgumentParser(description="导入科目四酒驾题 Word 题库")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--project", type=Path, default=Path.cwd())
    parser.add_argument("--apply", action="store_true")
    args = parser.parse_args()

    document = Document(args.docx)
    questions = parse_questions(document)
    answer_table = parse_answer_table(document)
    validate(questions, answer_table)

    report = {
        "question_count": len(questions),
        "single_choice_count": sum(question["type"] == "单选题" for question in questions),
        "multiple_choice_count": sum(question["type"] == "多选题" for question in questions),
        "source_image_count": len(document.inline_shapes),
        "answer_table_count": len(answer_table),
        "applied": False,
    }

    if args.apply:
        project = args.project.resolve()
        bank_path = project / "app" / "question-bank.json"
        records = build_records(questions, project / "public" / "question-images")
        bank = json.loads(bank_path.read_text(encoding="utf-8"))
        bank = [
            item
            for item in bank
            if item.get("section") not in {"A01", "C01"}
            and not item.get("code", "").startswith(("A01_", "C01_"))
        ]
        bank.extend(records)
        bank_path.write_text(
            json.dumps(bank, ensure_ascii=False, separators=(",", ":")),
            encoding="utf-8",
        )
        report.update(
            {
                "applied": True,
                "written_question_count": len(records),
                "written_image_count": len(records),
                "total_bank_count": len(bank),
                "codes": [record["code"] for record in records],
            }
        )

    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
