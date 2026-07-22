#!/usr/bin/env python3

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


QUESTION_RE = re.compile(r"^I03_(\d{2,3})\s+(.*)$")
OPTION_RE = re.compile(r"([ABCD])、(.*?)(?=\s+[ABCD]、|$)")
ANSWER_RE = re.compile(r"^【答案】([A-D]+)\s+【(单选题|多选题)】$")


def paragraph_image_parts(document, paragraph):
    parts = []
    for blip in paragraph._p.xpath(".//a:blip"):
        rel_id = blip.get(qn("r:embed"))
        if rel_id:
            parts.append(document.part.related_parts[rel_id])
    return parts


def parse_answer_table(document):
    if not document.tables:
        raise ValueError("文档中缺少答案速查表")

    answers = {}
    table = document.tables[-1]
    for row in table.rows[1:]:
        values = [cell.text.strip() for cell in row.cells]
        for index in range(0, len(values), 2):
            number = values[index]
            answer = values[index + 1] if index + 1 < len(values) else ""
            if number:
                answers[int(number)] = answer
    return answers


def parse_questions(document):
    questions = []
    current = None

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        match = QUESTION_RE.match(text)
        if match:
            if current:
                questions.append(current)
            number = int(match.group(1))
            current = {
                "number": number,
                "source_code": f"I03_{number:02d}",
                "title": match.group(2).strip(),
                "options": {},
                "answer": "",
                "type": "",
                "explanation": "",
                "detailedExplanation": "",
                "image_parts": [],
            }

        if current:
            current["image_parts"].extend(paragraph_image_parts(document, paragraph))

        if not current or not text or match:
            continue

        answer_match = ANSWER_RE.match(text)
        if answer_match:
            current["answer"] = answer_match.group(1)
            current["type"] = answer_match.group(2)
        elif text.startswith("【简析】"):
            current["explanation"] = text.removeprefix("【简析】").strip()
        elif text.startswith("【解析】"):
            current["detailedExplanation"] = text.removeprefix("【解析】").strip()
        elif text.startswith("A、"):
            current["options"] = {
                key: value.strip() for key, value in OPTION_RE.findall(text)
            }

    if current:
        questions.append(current)
    return questions


def validate(questions, answer_table):
    errors = []
    expected = list(range(1, 128))
    actual = [question["number"] for question in questions]
    if actual != expected:
        errors.append(f"题号不连续：expected 1..127, actual {actual[:5]}...{actual[-5:]}")

    for question in questions:
        number = question["number"]
        for field in ["title", "options", "answer", "type", "explanation", "detailedExplanation"]:
            if not question[field]:
                errors.append(f"第 {number} 题缺少 {field}")

        missing_answers = sorted(set(question["answer"]) - set(question["options"]))
        if missing_answers:
            errors.append(f"第 {number} 题答案不在选项中：{missing_answers}")

        table_answer = answer_table.get(number)
        if table_answer != question["answer"]:
            errors.append(
                f"第 {number} 题正文答案 {question['answer']} 与速查表 {table_answer} 不一致"
            )

    if len(answer_table) != 127:
        errors.append(f"答案速查表数量应为 127，实际为 {len(answer_table)}")

    if errors:
        raise ValueError("\n".join(errors))


def build_bank_records(questions, image_dir):
    records = []
    image_count = 0
    repeated_image_questions = []

    for question in questions:
        number = question["number"]
        code = f"L01_{number:03d}"
        image = ""
        if question["image_parts"]:
            image_count += 1
            if len(question["image_parts"]) > 1:
                repeated_image_questions.append(code)
            part = question["image_parts"][0]
            suffix = Path(str(part.partname)).suffix.lower() or ".png"
            filename = f"l01-{number:03d}{suffix}"
            (image_dir / filename).write_bytes(part.blob)
            image = f"/question-images/{filename}"

        options = {key: question["options"].get(key, "") for key in "ABCD"}
        records.append(
            {
                "id": f"subject4-lights-{code}",
                "code": code,
                "title": question["title"],
                "image": image,
                "options": options,
                "type": question["type"],
                "category": "科四灯光类-L",
                "section": "L01",
                "difficulty": "基础",
                "answer": question["answer"],
                "explanation": question["explanation"],
                "detailedExplanation": question["detailedExplanation"],
                "tags": ["科四灯光", "L01"],
                "status": "已发布",
                "updatedAt": "2026-07-22",
            }
        )

    return records, image_count, repeated_image_questions


def main():
    parser = argparse.ArgumentParser(description="导入科目四灯光题 Word 题库")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--project", type=Path, default=Path.cwd())
    parser.add_argument("--apply", action="store_true")
    args = parser.parse_args()

    document = Document(args.docx)
    questions = parse_questions(document)
    answer_table = parse_answer_table(document)
    validate(questions, answer_table)

    image_questions = [question for question in questions if question["image_parts"]]
    report = {
        "question_count": len(questions),
        "single_choice_count": sum(question["type"] == "单选题" for question in questions),
        "multiple_choice_count": sum(question["type"] == "多选题" for question in questions),
        "image_question_count": len(image_questions),
        "image_reference_count": sum(len(question["image_parts"]) for question in questions),
        "image_question_codes": [f"L01_{question['number']:03d}" for question in image_questions],
        "answer_table_count": len(answer_table),
        "applied": False,
    }

    if args.apply:
        project = args.project.resolve()
        bank_path = project / "app" / "question-bank.json"
        image_dir = project / "public" / "question-images"
        image_dir.mkdir(parents=True, exist_ok=True)

        records, image_count, repeated = build_bank_records(questions, image_dir)
        bank = json.loads(bank_path.read_text(encoding="utf-8"))
        bank = [
            item
            for item in bank
            if item.get("section") != "L01" and not item.get("code", "").startswith("L01_")
        ]
        bank.extend(records)
        bank_path.write_text(
            json.dumps(bank, ensure_ascii=False, separators=(",", ":")),
            encoding="utf-8",
        )
        report.update(
            {
                "applied": True,
                "written_question_count": len(records),
                "written_image_count": image_count,
                "multiple_image_questions": repeated,
                "total_bank_count": len(bank),
            }
        )

    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
