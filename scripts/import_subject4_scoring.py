#!/usr/bin/env python3

import argparse
import json
import re
from pathlib import Path

from docx import Document


QUESTION_RE = re.compile(r"^L02_(\d{2,3})\s+(.*)$")
OPTION_RE = re.compile(r"([ABCD])、(.*?)(?=\s+[ABCD]、|$)")
ANSWER_RE = re.compile(r"^【答案】([A-D]+)\s+【(单选题|多选题)】$")


def parse_answer_table(document):
    if not document.tables:
        raise ValueError("文档中缺少答案速查表")
    answers = {}
    for row in document.tables[-1].rows[1:]:
        values = [cell.text.strip() for cell in row.cells]
        for index in range(0, len(values), 2):
            number = values[index]
            answer = values[index + 1] if index + 1 < len(values) else ""
            if number:
                answers[int(number)] = answer
    return answers


def parse_questions(document):
    questions = []
    current = None
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        match = QUESTION_RE.match(text)
        if match:
            if current:
                questions.append(current)
            current = {
                "number": int(match.group(1)),
                "title": match.group(2).strip(),
                "options": {},
                "answer": "",
                "type": "",
                "explanation": "",
                "detailedExplanation": "",
            }
            continue
        if not current or not text:
            continue
        answer_match = ANSWER_RE.match(text)
        if answer_match:
            current["answer"] = answer_match.group(1)
            current["type"] = answer_match.group(2)
        elif text.startswith("【简析】"):
            current["explanation"] = text.removeprefix("【简析】").strip()
        elif text.startswith("【解析】"):
            current["detailedExplanation"] = text.removeprefix("【解析】").strip()
        elif text.startswith("A、"):
            current["options"] = {
                key: value.strip() for key, value in OPTION_RE.findall(text)
            }
    if current:
        questions.append(current)
    return questions


def validate(questions, answer_table):
    errors = []
    if [question["number"] for question in questions] != [1]:
        errors.append("题号应为 L02_01")
    for question in questions:
        for field in ["title", "options", "answer", "type", "explanation", "detailedExplanation"]:
            if not question[field]:
                errors.append(f"第 {question['number']} 题缺少 {field}")
        if answer_table.get(question["number"]) != question["answer"]:
            errors.append(
                f"第 {question['number']} 题正文答案 {question['answer']} "
                f"与速查表 {answer_table.get(question['number'])} 不一致"
            )
        missing = sorted(set(question["answer"]) - set(question["options"]))
        if missing:
            errors.append(f"第 {question['number']} 题答案不在选项中：{missing}")
    if len(answer_table) != 1:
        errors.append(f"答案速查表数量应为 1，实际为 {len(answer_table)}")
    if errors:
        raise ValueError("\n".join(errors))


def build_prompt(question):
    options = "；".join(f"{key}、{value}" for key, value in question["options"].items())
    return f"""Use case: photorealistic-natural, scientific-educational
Asset type: Chinese driving-theory learning cover image, landscape 16:9, readable at 375px mobile width
Primary request: Create one coherent photorealistic teaching scene for this question: {question['title']}
Question options: {options}
Correct answer: {question['answer']}
Teaching point: all four listed situations are valid reasons to remove an incorrectly or exceptionally recorded traffic violation.
Scene/backdrop: a realistic Chinese traffic-police evidence review center. One reviewer is examining a large traffic-monitoring screen containing four clearly separated CCTV evidence cards within the same workstation scene.
Required visual facts: card 1 shows conflicting traffic-signal indications; card 2 shows a private vehicle moving aside to rescue a person in danger or avoid an emergency; card 3 shows a verified stolen vehicle case; card 4 shows a cloned or forged plate causing the lawful vehicle to be recorded by mistake. Make all four cases understandable without graphic harm.
Composition/framing: medium-wide documentary view over the reviewer's shoulder; the four evidence cards are large, balanced and readable; a single clear approved-removal state connects all four cards.
Style/medium: photorealistic documentary photography, realistic Chinese roads, cameras, vehicles and control-room equipment, neutral professional lighting.
Critical accuracy: do not depict ordinary valid violations being excused. The lawful vehicle in the cloned-plate case must be visibly distinct from the offending vehicle. The rescue/emergency card must clearly show necessity rather than convenience.
Constraints: no police insignia, agency logo, brand, watermark, gore, injury detail, unrelated sign, invented statute number or dense interface text. Avoid tiny or garbled Chinese text; prefer clear visual storytelling and simple universally understood approval/deletion iconography.
"""


def build_records(questions, image_dir):
    records = []
    for question in questions:
        number = question["number"]
        code = f"C03_{number + 5:02d}"
        image_name = f"c03-{number + 5:02d}-v2.webp"
        if not (image_dir / image_name).is_file():
            raise FileNotFoundError(f"缺少重制题图：{image_name}")
        records.append(
            {
                "id": f"subject4-scoring-{code}",
                "code": code,
                "title": question["title"],
                "image": f"/question-images/{image_name}",
                "options": {key: question["options"].get(key, "") for key in "ABCD"},
                "type": question["type"],
                "category": "科四场景类-C",
                "section": "C03",
                "difficulty": "基础",
                "answer": question["answer"],
                "explanation": question["explanation"],
                "detailedExplanation": question["detailedExplanation"],
                "tags": ["科四记分", "C03"],
                "status": "已发布",
                "updatedAt": "2026-07-22",
            }
        )
    return records


def main():
    parser = argparse.ArgumentParser(description="导入科目四记分题 Word 题库")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--project", type=Path, default=Path.cwd())
    parser.add_argument("--prompt-manifest", type=Path)
    parser.add_argument("--apply", action="store_true")
    args = parser.parse_args()

    document = Document(args.docx)
    questions = parse_questions(document)
    answer_table = parse_answer_table(document)
    validate(questions, answer_table)

    if args.prompt_manifest:
        manifest = [
            {
                "number": question["number"],
                "code": f"C03_{question['number'] + 5:02d}",
                "prompt": build_prompt(question),
                "output_image": f"public/question-images/c03-{question['number'] + 5:02d}-v2.webp",
            }
            for question in questions
        ]
        args.prompt_manifest.parent.mkdir(parents=True, exist_ok=True)
        args.prompt_manifest.write_text(
            json.dumps(manifest, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    report = {
        "question_count": len(questions),
        "single_choice_count": sum(question["type"] == "单选题" for question in questions),
        "multiple_choice_count": sum(question["type"] == "多选题" for question in questions),
        "answer_table_count": len(answer_table),
        "prompt_manifest": str(args.prompt_manifest) if args.prompt_manifest else None,
        "applied": False,
    }

    if args.apply:
        project = args.project.resolve()
        bank_path = project / "app" / "question-bank.json"
        records = build_records(questions, project / "public" / "question-images")
        bank = json.loads(bank_path.read_text(encoding="utf-8"))
        bank = [
            item
            for item in bank
            if item.get("id") not in {"subject4-scoring-C04_01", "subject4-scoring-C03_06"}
            and item.get("code") not in {"C04_01", "C03_06", "L02_01"}
        ]
        bank.extend(records)
        bank_path.write_text(
            json.dumps(bank, ensure_ascii=False, separators=(",", ":")),
            encoding="utf-8",
        )
        report.update(
            {
                "applied": True,
                "written_question_count": len(records),
                "written_image_count": len(records),
                "total_bank_count": len(bank),
                "codes": [record["code"] for record in records],
            }
        )

    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
