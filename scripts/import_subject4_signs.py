#!/usr/bin/env python3

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image
from io import BytesIO


QUESTION_RE = re.compile(r"^I05_(\d{2,3})\s+(.*)$")
OPTION_RE = re.compile(r"([ABCD])、(.*?)(?=\s+[ABCD]、|$)")
ANSWER_RE = re.compile(r"^【答案】([A-D]+)\s+【(单选题|多选题)】$")


def parse_answer_table(document):
    if not document.tables:
        raise ValueError("文档中缺少答案速查表")
    answers = {}
    for row in document.tables[-1].rows[1:]:
        values = [cell.text.strip() for cell in row.cells]
        for index in range(0, len(values), 2):
            number = values[index]
            answer = values[index + 1] if index + 1 < len(values) else ""
            if number:
                answers[int(number)] = answer
    return answers


def paragraph_image_parts(document, paragraph):
    parts = []
    for element in paragraph._p.iter():
        if element.tag == qn("a:blip"):
            relationship_id = element.get(qn("r:embed"))
            if relationship_id:
                parts.append(document.part.rels[relationship_id].target_part)
    return parts


def parse_questions(document):
    questions = []
    current = None

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        question_match = QUESTION_RE.match(text)
        if question_match:
            if current:
                questions.append(current)
            current = {
                "number": int(question_match.group(1)),
                "title": question_match.group(2).strip(),
                "options": {},
                "answer": "",
                "type": "",
                "explanation": "",
                "detailedExplanation": "",
                "source_image_blob": None,
            }
            continue

        if not current:
            continue

        image_parts = paragraph_image_parts(document, paragraph)
        if image_parts and current["source_image_blob"] is None:
            current["source_image_blob"] = image_parts[0].blob

        if not text:
            continue
        answer_match = ANSWER_RE.match(text)
        if answer_match:
            current["answer"] = answer_match.group(1)
            current["type"] = answer_match.group(2)
        elif text.startswith("【简析】"):
            current["explanation"] = text.removeprefix("【简析】").strip()
        elif text.startswith("【解析】"):
            current["detailedExplanation"] = text.removeprefix("【解析】").strip()
        elif text.startswith("A、"):
            current["options"] = {
                key: value.strip() for key, value in OPTION_RE.findall(text)
            }

    if current:
        questions.append(current)
    return questions


def validate(questions, answer_table):
    errors = []
    expected = list(range(1, 221))
    actual = [question["number"] for question in questions]
    if actual != expected:
        errors.append(f"题号不连续：expected 1..220, actual {actual}")

    for question in questions:
        number = question["number"]
        for field in ["title", "options", "answer", "type", "explanation", "detailedExplanation"]:
            if not question[field]:
                errors.append(f"第 {number} 题缺少 {field}")
        missing_answers = sorted(set(question["answer"]) - set(question["options"]))
        if missing_answers:
            errors.append(f"第 {number} 题答案不在选项中：{missing_answers}")
        if answer_table.get(number) != question["answer"]:
            errors.append(
                f"第 {number} 题正文答案 {question['answer']} 与速查表 {answer_table.get(number)} 不一致"
            )

    if len(answer_table) != 220:
        errors.append(f"答案速查表数量应为 220，实际为 {len(answer_table)}")
    if errors:
        raise ValueError("\n".join(errors))


def extract_source_images(questions, output_dir):
    output_dir.mkdir(parents=True, exist_ok=True)
    count = 0
    for question in questions:
        blob = question["source_image_blob"]
        if blob is None:
            continue
        image = Image.open(BytesIO(blob)).convert("RGB")
        path = output_dir / f"i05-{question['number']:02d}-source.png"
        image.save(path, "PNG")
        count += 1
    return count


def build_prompt_manifest(questions, source_dir):
    manifest = []
    for question in questions:
        number = question["number"]
        source_path = source_dir / f"i05-{number:02d}-source.png"
        options = "；".join(f"{key}、{value}" for key, value in question["options"].items())
        reference_instruction = (
            "Input image role: authoritative traffic-layout and traffic-sign reference only. "
            "Copy the sign face literally from the reference, including every line, arrow, number and Chinese character; "
            "do not substitute a semantically related sign. Do not retain the cartoon rendering style."
            if source_path.is_file()
            else "No input image is available. Build the scene strictly from the question, correct answer and authoritative explanation."
        )
        prompt = f"""Use case: photorealistic-natural, scientific-educational
Asset type: Chinese driving-theory learning cover image, landscape 16:9, readable at 375px mobile width
Primary request: Completely rebuild the supplied old traffic-theory illustration as a brand-new photorealistic Chinese road scene. The image must accurately communicate this question: {question['title']}
Question options: {options}
Correct answer: {question['answer']}
Teaching point: {question['explanation']}
Authoritative explanation: {question['detailedExplanation']}
{reference_instruction}
Critical accuracy: preserve the exact sign category, sign shape, border color, background color, symbol, arrow direction, number, Chinese character, lane direction, vehicle type, vehicle position, people and driving action required by the question and reference. A visually similar or related sign is not acceptable. If the reference sign has no text, add no text. If it contains a number or Chinese character, render it exactly and do not invent extra text. The road geometry and driving viewpoint must be physically plausible in China.
Style/medium: photorealistic documentary driving scene, natural daylight or source-appropriate lighting, realistic vehicle geometry, driver-eye viewpoint when useful, crisp sign visibility.
Composition/framing: one coherent scene, important sign and action large enough to inspect on mobile, no montage unless multiple facts are essential.
Constraints: safe educational presentation; no injuries or gore; no unrelated vehicles, signs, arrows, numbers or text; no brands, logos, watermark, cartoon, illustration or video-game look.
"""
        manifest.append(
            {
                "number": number,
                "code": f"I05_{number:02d}",
                "prompt": prompt,
                "source_image": str(source_path) if source_path.is_file() else None,
                "output_image": f"public/question-images/i05-{number:02d}-v2.webp",
            }
        )
    return manifest


def build_records(questions, image_dir):
    records = []
    for question in questions:
        number = question["number"]
        code = f"I05_{number:02d}"
        image_name = f"i05-{number:02d}-v2.webp"
        if not (image_dir / image_name).is_file():
            raise FileNotFoundError(f"缺少重制题图：{image_name}")
        records.append(
            {
                "id": f"subject4-signs-{code}",
                "code": code,
                "title": question["title"],
                "image": f"/question-images/{image_name}",
                "options": {key: question["options"].get(key, "") for key in "ABCD"},
                "type": question["type"],
                "category": "科四标识类-I",
                "section": "I05",
                "difficulty": "基础",
                "answer": question["answer"],
                "explanation": question["explanation"],
                "detailedExplanation": question["detailedExplanation"],
                "tags": ["科四标识", "I05"],
                "status": "已发布",
                "updatedAt": "2026-07-22",
            }
        )
    return records


def main():
    parser = argparse.ArgumentParser(description="导入科目四标志题 Word 题库")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--project", type=Path, default=Path.cwd())
    parser.add_argument("--extract-source-images", type=Path)
    parser.add_argument("--prompt-manifest", type=Path)
    parser.add_argument("--apply", action="store_true")
    args = parser.parse_args()

    document = Document(args.docx)
    questions = parse_questions(document)
    answer_table = parse_answer_table(document)
    validate(questions, answer_table)

    source_dir = args.extract_source_images or Path("/private/tmp/subject4_sign_sources")
    extracted = None
    if args.extract_source_images:
        extracted = extract_source_images(questions, source_dir)
    if args.prompt_manifest:
        manifest = build_prompt_manifest(questions, source_dir)
        args.prompt_manifest.parent.mkdir(parents=True, exist_ok=True)
        args.prompt_manifest.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    report = {
        "question_count": len(questions),
        "single_choice_count": sum(question["type"] == "单选题" for question in questions),
        "multiple_choice_count": sum(question["type"] == "多选题" for question in questions),
        "source_image_reference_count": sum(question["source_image_blob"] is not None for question in questions),
        "answer_table_count": len(answer_table),
        "extracted_source_image_count": extracted,
        "prompt_manifest": str(args.prompt_manifest) if args.prompt_manifest else None,
        "applied": False,
    }

    if args.apply:
        project = args.project.resolve()
        bank_path = project / "app" / "question-bank.json"
        records = build_records(questions, project / "public" / "question-images")
        bank = json.loads(bank_path.read_text(encoding="utf-8"))
        bank = [
            item
            for item in bank
            if item.get("section") != "I05" and not item.get("code", "").startswith("I05_")
        ]
        bank.extend(records)
        bank_path.write_text(
            json.dumps(bank, ensure_ascii=False, separators=(",", ":")),
            encoding="utf-8",
        )
        report.update(
            {
                "applied": True,
                "written_question_count": len(records),
                "written_image_count": len(records),
                "total_bank_count": len(bank),
                "codes": [records[0]["code"], records[-1]["code"]],
            }
        )

    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
